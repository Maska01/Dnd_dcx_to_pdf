import os
import re
from dataclasses import dataclass, field

from docx import Document
from PIL import Image as PILImage
from reportlab.lib.styles import ParagraphStyle
from reportlab.platypus import BaseDocTemplate, Frame, ListFlowable, ListItem, PageBreak, PageTemplate, Paragraph, Spacer, Image
from reportlab.platypus.tableofcontents import TableOfContents

import configuracion_pdf as cfg
from procesamiento_word import (
    agregar_imagenes_a_bloque,
    decorar_consejo_dm_html,
    decorar_info_adicional_html,
    es_consejo_dm,
    es_fin_bloque_manual,
    es_info_adicional,
    es_lista_parrafo,
    es_inicio_bloque_aliado,
    es_inicio_bloque_consejo,
    es_inicio_bloque_cita,
    es_inicio_bloque_enemigo,
    es_inicio_bloque_info,
    es_inicio_bloque_npc,
    es_inicio_bloque_objeto,
    es_inicio_bloque_tesoro,
    estilo_para_parrafo,
    extraer_imagenes_de_parrafo,
    item_caja_desde_parrafo,
    item_caja_plano,
    item_caja_tabla,
    iterar_elementos_documento,
    parrafo_a_html,
    _extraer_bloque_prefijo_embebido,
    _parrafo_es_solo_salto_de_pagina,
    _parrafo_tiene_salto_de_pagina_previo,
)
from renderizado_cajas import (
    crear_flujos_imagenes,
    decorar_aliado_en_html,
    decorar_enemigo_en_html,
    decorar_npc_en_html,
    decorar_objeto_en_html,
    decorar_premio_en_html,
    decorar_tesoro_en_html,
    renderizar_caja,
    tabla_docx_a_flujo,
)


class DocumentoConIndice(BaseDocTemplate):
    def __init__(self, filename, **kw):
        super().__init__(filename, **kw)
        frame = Frame(self.leftMargin, self.bottomMargin, self.width, self.height, id="normal")
        self.addPageTemplates([PageTemplate(id="Todo", frames=frame, onPage=self._dibujar_fondo_pagina)])
        self._contador_marcadores = 0

    def beforeDocument(self):
        super().beforeDocument()
        self._contador_marcadores = 0

    def _dibujar_fondo_pagina(self, canvas, doc):
        canvas.saveState()
        canvas.setFillColor(cfg.COLOR_FONDO_PAGINA)
        canvas.rect(0, 0, doc.pagesize[0], doc.pagesize[1], stroke=0, fill=1)
        canvas.restoreState()

    def _crear_marcador(self, texto):
        self._contador_marcadores += 1
        texto_limpio = re.sub(r"\s+", "-", texto.strip())
        texto_limpio = re.sub(r"[^\w\-]", "", texto_limpio, flags=re.UNICODE)
        texto_limpio = texto_limpio[:60] or "seccion"
        return f"marcador-{self._contador_marcadores}-{texto_limpio}"

    def afterFlowable(self, flowable):
        if isinstance(flowable, Paragraph):
            estilo = flowable.style.name
            texto = flowable.getPlainText()
            if estilo == "TituloIndice":
                marcador = "indice"
                self.canv.bookmarkPage(marcador)
                self.canv.addOutlineEntry(texto or "Índice", marcador, level=0, closed=False)
            elif estilo == "H1":
                marcador = self._crear_marcador(texto)
                self.canv.bookmarkPage(marcador)
                self.canv.addOutlineEntry(texto, marcador, level=0, closed=False)
                self.notify("TOCEntry", (0, texto, self.page, marcador))
            elif estilo == "H2":
                marcador = self._crear_marcador(texto)
                self.canv.bookmarkPage(marcador)
                self.canv.addOutlineEntry(texto, marcador, level=1, closed=False)
                self.notify("TOCEntry", (1, texto, self.page, marcador))
            elif estilo == "H3":
                marcador = self._crear_marcador(texto)
                self.canv.bookmarkPage(marcador)
                self.canv.addOutlineEntry(texto, marcador, level=2, closed=False)


@dataclass
class EstadoConstruccionPdf:
    historia: list
    estilos: dict
    ancho_util: float
    items_lista_actual: list = field(default_factory=list)
    citas_pendientes: list = field(default_factory=list)
    vacios_desde_ultima_cita: int = 0
    max_vacios_fusion_cita: int = 2
    consejos_pendientes: list = field(default_factory=list)
    modo_consejos: str | None = None
    vacios_desde_ultimo_consejo: int = 0
    max_vacios_fusion_consejo: int = 2
    infos_pendientes: list = field(default_factory=list)
    modo_infos: str | None = None
    vacios_desde_ultima_info: int = 0
    max_vacios_fusion_info: int = 2
    bloque_consejo_manual: list = field(default_factory=list)
    dentro_consejo_manual: bool = False
    bloque_cita_manual: list = field(default_factory=list)
    dentro_cita_manual: bool = False
    bloque_npc_manual: list = field(default_factory=list)
    dentro_npc_manual: bool = False
    bloque_enemigo_manual: list = field(default_factory=list)
    dentro_enemigo_manual: bool = False
    bloque_aliado_manual: list = field(default_factory=list)
    dentro_aliado_manual: bool = False
    bloque_tesoro_manual: list = field(default_factory=list)
    dentro_tesoro_manual: bool = False
    titulo_tesoro_manual: str = "Tesoro"
    bloque_objeto_manual: list = field(default_factory=list)
    dentro_objeto_manual: bool = False
    bloque_info_manual: list = field(default_factory=list)
    dentro_info: bool = False

    def agregar_caja_a_historia(self, partes, estilo, decorador=None):
        caja = renderizar_caja(partes, self.estilos[estilo], self.ancho_util, decorador=decorador)
        if caja:
            self.historia.extend(caja)

    def agregar_contenido_a_bloque_activo(self, item_caja, imagenes):
        destinos = [
            (self.dentro_info, self.bloque_info_manual),
            (self.dentro_consejo_manual, self.bloque_consejo_manual),
            (self.dentro_cita_manual, self.bloque_cita_manual),
            (self.dentro_npc_manual, self.bloque_npc_manual),
            (self.dentro_enemigo_manual, self.bloque_enemigo_manual),
            (self.dentro_aliado_manual, self.bloque_aliado_manual),
            (self.dentro_tesoro_manual, self.bloque_tesoro_manual),
            (self.dentro_objeto_manual, self.bloque_objeto_manual),
        ]
        for activo, bloque in destinos:
            if activo:
                if item_caja is not None:
                    bloque.append(item_caja)
                if imagenes:
                    agregar_imagenes_a_bloque(bloque, imagenes)
                return True
        return False

    def agregar_salto_a_bloque_activo(self):
        for activo, bloque in [
            (self.dentro_consejo_manual, self.bloque_consejo_manual),
            (self.dentro_cita_manual, self.bloque_cita_manual),
            (self.dentro_npc_manual, self.bloque_npc_manual),
            (self.dentro_enemigo_manual, self.bloque_enemigo_manual),
            (self.dentro_aliado_manual, self.bloque_aliado_manual),
            (self.dentro_tesoro_manual, self.bloque_tesoro_manual),
            (self.dentro_objeto_manual, self.bloque_objeto_manual),
            (self.dentro_info, self.bloque_info_manual),
        ]:
            if activo:
                bloque.append(None)
                return True
        return False

    def vaciar_lista(self):
        if self.items_lista_actual:
            self.historia.append(ListFlowable([ListItem(Paragraph(texto, self.estilos["Cuerpo"]), leftIndent=12) for texto in self.items_lista_actual], bulletType="bullet", start="•"))
            self.historia.append(Spacer(1, 0))
            self.items_lista_actual.clear()

    def vaciar_citas(self):
        if self.citas_pendientes:
            self.agregar_caja_a_historia(self.citas_pendientes, "CajaCita")
            self.citas_pendientes.clear()

    def vaciar_consejos(self):
        if self.consejos_pendientes:
            self.agregar_caja_a_historia(self.consejos_pendientes, "CajaConsejoDm", decorar_consejo_dm_html)
            self.consejos_pendientes.clear()
        self.vacios_desde_ultimo_consejo = 0
        self.modo_consejos = None

    def vaciar_infos(self):
        if self.infos_pendientes:
            self.agregar_caja_a_historia(self.infos_pendientes, "CajaInfoAdicional", decorar_info_adicional_html)
            self.infos_pendientes.clear()
        self.vacios_desde_ultima_info = 0
        self.modo_infos = None

    def emitir_consejo_manual(self):
        if self.bloque_consejo_manual:
            self.agregar_caja_a_historia(self.bloque_consejo_manual, "CajaConsejoDm", decorar_consejo_dm_html)
            self.bloque_consejo_manual.clear()
        self.dentro_consejo_manual = False

    def emitir_cita_manual(self):
        if self.bloque_cita_manual:
            self.agregar_caja_a_historia(self.bloque_cita_manual, "CajaCita")
            self.bloque_cita_manual.clear()
        self.dentro_cita_manual = False

    def emitir_npc_manual(self):
        if self.bloque_npc_manual:
            self.agregar_caja_a_historia(self.bloque_npc_manual, "CajaNpc", decorar_npc_en_html)
            self.bloque_npc_manual.clear()
        self.dentro_npc_manual = False

    def emitir_enemigo_manual(self):
        if self.bloque_enemigo_manual:
            self.agregar_caja_a_historia(self.bloque_enemigo_manual, "CajaEnemigo", decorar_enemigo_en_html)
            self.bloque_enemigo_manual.clear()
        self.dentro_enemigo_manual = False

    def emitir_aliado_manual(self):
        if self.bloque_aliado_manual:
            self.agregar_caja_a_historia(self.bloque_aliado_manual, "CajaAliado", decorar_aliado_en_html)
            self.bloque_aliado_manual.clear()
        self.dentro_aliado_manual = False

    def emitir_tesoro_manual(self):
        if self.bloque_tesoro_manual:
            decorador = decorar_premio_en_html if self.titulo_tesoro_manual == "Premio" else decorar_tesoro_en_html
            self.agregar_caja_a_historia(self.bloque_tesoro_manual, "CajaTesoro", decorador)
            self.bloque_tesoro_manual.clear()
        self.dentro_tesoro_manual = False
        self.titulo_tesoro_manual = "Tesoro"

    def emitir_objeto_manual(self):
        if self.bloque_objeto_manual:
            self.agregar_caja_a_historia(self.bloque_objeto_manual, "CajaObjeto", decorar_objeto_en_html)
            self.bloque_objeto_manual.clear()
        self.dentro_objeto_manual = False

    def emitir_info_adicional(self):
        if self.bloque_info_manual:
            self.agregar_caja_a_historia(self.bloque_info_manual, "CajaInfoAdicional", decorar_info_adicional_html)
            self.bloque_info_manual.clear()
        self.dentro_info = False

    def insertar_salto_de_pagina(self):
        self.vaciar_consejos()
        self.vaciar_infos()
        self.emitir_info_adicional()
        self.emitir_consejo_manual()
        self.emitir_cita_manual()
        self.emitir_npc_manual()
        self.emitir_enemigo_manual()
        self.emitir_aliado_manual()
        self.emitir_tesoro_manual()
        self.emitir_objeto_manual()
        self.vaciar_citas()
        self.vaciar_lista()
        if not self.historia or not isinstance(self.historia[-1], PageBreak):
            self.historia.append(PageBreak())

    def cerrar_bloques_para_contenido_general(self):
        self.vaciar_consejos()
        self.vaciar_infos()
        self.emitir_info_adicional()
        self.emitir_consejo_manual()
        self.emitir_cita_manual()
        self.emitir_npc_manual()
        self.emitir_enemigo_manual()
        self.emitir_aliado_manual()
        self.emitir_tesoro_manual()
        self.emitir_objeto_manual()
        self.vaciar_citas()
        self.vaciar_lista()

    def cerrar_todo_al_final(self):
        self.emitir_consejo_manual()
        self.emitir_cita_manual()
        self.emitir_npc_manual()
        self.emitir_enemigo_manual()
        self.emitir_aliado_manual()
        self.emitir_tesoro_manual()
        self.emitir_objeto_manual()
        self.emitir_info_adicional()
        self.vaciar_consejos()
        self.vaciar_infos()
        self.vaciar_citas()
        self.vaciar_lista()

    def procesar_fin_bloque_manual(self, texto_plano):
        if es_fin_bloque_manual(texto_plano) and self.dentro_info:
            self.emitir_info_adicional()
            return True
        if es_fin_bloque_manual(texto_plano) and self.dentro_consejo_manual:
            self.emitir_consejo_manual()
            return True
        if es_fin_bloque_manual(texto_plano) and self.dentro_cita_manual:
            self.emitir_cita_manual()
            return True
        if es_fin_bloque_manual(texto_plano) and self.dentro_npc_manual:
            self.emitir_npc_manual()
            return True
        if es_fin_bloque_manual(texto_plano) and self.dentro_enemigo_manual:
            self.emitir_enemigo_manual()
            return True
        if es_fin_bloque_manual(texto_plano) and self.dentro_aliado_manual:
            self.emitir_aliado_manual()
            return True
        if es_fin_bloque_manual(texto_plano) and self.dentro_tesoro_manual:
            self.emitir_tesoro_manual()
            return True
        if es_fin_bloque_manual(texto_plano) and self.dentro_objeto_manual:
            self.emitir_objeto_manual()
            return True
        return False

    def procesar_apertura_bloque_manual(self, texto_plano):
        if es_inicio_bloque_info(texto_plano):
            self.vaciar_consejos(); self.vaciar_infos(); self.emitir_consejo_manual(); self.emitir_cita_manual(); self.emitir_npc_manual(); self.emitir_enemigo_manual(); self.emitir_aliado_manual(); self.vaciar_citas(); self.vaciar_lista(); self.dentro_info = True; self.bloque_info_manual.clear(); return True
        if es_inicio_bloque_consejo(texto_plano):
            self.vaciar_consejos(); self.vaciar_infos(); self.emitir_info_adicional(); self.emitir_cita_manual(); self.emitir_npc_manual(); self.emitir_enemigo_manual(); self.emitir_aliado_manual(); self.vaciar_citas(); self.vaciar_lista(); self.dentro_consejo_manual = True; self.bloque_consejo_manual.clear(); return True
        if es_inicio_bloque_cita(texto_plano):
            self.vaciar_consejos(); self.vaciar_infos(); self.emitir_info_adicional(); self.emitir_consejo_manual(); self.emitir_npc_manual(); self.emitir_enemigo_manual(); self.emitir_aliado_manual(); self.vaciar_citas(); self.vaciar_lista(); self.dentro_cita_manual = True; self.bloque_cita_manual.clear(); return True
        if es_inicio_bloque_npc(texto_plano):
            self.vaciar_consejos(); self.vaciar_infos(); self.emitir_info_adicional(); self.emitir_consejo_manual(); self.emitir_cita_manual(); self.emitir_enemigo_manual(); self.emitir_aliado_manual(); self.vaciar_citas(); self.vaciar_lista(); self.dentro_npc_manual = True; self.bloque_npc_manual.clear(); return True
        if es_inicio_bloque_enemigo(texto_plano):
            self.vaciar_consejos(); self.vaciar_infos(); self.emitir_info_adicional(); self.emitir_consejo_manual(); self.emitir_cita_manual(); self.emitir_npc_manual(); self.emitir_aliado_manual(); self.vaciar_citas(); self.vaciar_lista(); self.dentro_enemigo_manual = True; self.bloque_enemigo_manual.clear(); return True
        if es_inicio_bloque_aliado(texto_plano):
            self.vaciar_consejos(); self.vaciar_infos(); self.emitir_info_adicional(); self.emitir_consejo_manual(); self.emitir_cita_manual(); self.emitir_npc_manual(); self.emitir_enemigo_manual(); self.vaciar_citas(); self.vaciar_lista(); self.dentro_aliado_manual = True; self.bloque_aliado_manual.clear(); return True
        if es_inicio_bloque_tesoro(texto_plano):
            self.vaciar_consejos(); self.vaciar_infos(); self.emitir_info_adicional(); self.emitir_consejo_manual(); self.emitir_cita_manual(); self.emitir_npc_manual(); self.emitir_enemigo_manual(); self.emitir_aliado_manual(); self.emitir_objeto_manual(); self.vaciar_citas(); self.vaciar_lista(); self.dentro_tesoro_manual = True; self.titulo_tesoro_manual = "Premio" if "premio" in (texto_plano or "").lower() else "Tesoro"; self.bloque_tesoro_manual.clear(); return True
        if es_inicio_bloque_objeto(texto_plano):
            self.vaciar_consejos(); self.vaciar_infos(); self.emitir_info_adicional(); self.emitir_consejo_manual(); self.emitir_cita_manual(); self.emitir_npc_manual(); self.emitir_enemigo_manual(); self.emitir_aliado_manual(); self.emitir_tesoro_manual(); self.vaciar_citas(); self.vaciar_lista(); self.dentro_objeto_manual = True; self.bloque_objeto_manual.clear(); return True
        return False


def construir_pdf(ruta_docx, ruta_pdf, titulo=None, autor=None, subtitulo=None, imagen_portada=None):
    documento_word = Document(ruta_docx)
    estilos = cfg.construir_estilos()
    documento_pdf = DocumentoConIndice(str(ruta_pdf), pagesize=cfg.TAMANO_PAGINA, leftMargin=cfg.MARGEN, rightMargin=cfg.MARGEN, topMargin=cfg.MARGEN, bottomMargin=cfg.MARGEN, title=titulo or "Aventura", author=autor or "")
    ancho_util = cfg.TAMANO_PAGINA[0] - 2 * cfg.MARGEN
    historia = []
    if titulo or subtitulo or autor or imagen_portada:
        if imagen_portada and os.path.exists(imagen_portada):
            try:
                with PILImage.open(imagen_portada) as imagen:
                    ancho, alto = imagen.size
                ratio = (alto / ancho) if ancho else 1
                imagen_portada_flujo = Image(imagen_portada, width=ancho_util, height=ancho_util * ratio)
                imagen_portada_flujo.hAlign = "CENTER"
                historia.extend([Spacer(1, 1 * cfg.cm), imagen_portada_flujo, Spacer(1, 1 * cfg.cm)])
            except Exception as error:
                print(f"⚠️  No se pudo cargar la imagen de portada: {error}")
                historia.append(Spacer(1, 6 * cfg.cm))
        else:
            if imagen_portada:
                print(f"ℹ️  Imagen de portada no encontrada en: {imagen_portada} (se omite). Cambia la ruta cuando tengas la imagen.")
            historia.append(Spacer(1, 6 * cfg.cm))
        if titulo:
            historia.append(Paragraph(titulo, estilos["PortadaTitulo"]))
        if subtitulo:
            historia.append(Paragraph(subtitulo, estilos["PortadaSubtitulo"]))
        if autor:
            historia.extend([Spacer(1, 1 * cfg.cm), Paragraph(f"por <b>{autor}</b>", estilos["PortadaAutor"])])
        historia.append(PageBreak())
    historia.append(Paragraph("Índice", estilos["TituloIndice"]))
    indice = TableOfContents()
    indice.levelStyles = [ParagraphStyle(name="IndiceNivel1", fontName=cfg.FUENTE_TITULO, fontSize=12, leading=18, textColor=cfg.COLOR_PRIMARIO, leftIndent=0), ParagraphStyle(name="IndiceNivel2", fontName=cfg.FUENTE_TEXTO, fontSize=11, leading=16, textColor=cfg.COLOR_SECUNDARIO, leftIndent=18)]
    historia.extend([indice, PageBreak()])
    estado = EstadoConstruccionPdf(historia=historia, estilos=estilos, ancho_util=ancho_util)
    estilo_tabla_general = ParagraphStyle(name="TablaGeneral", parent=estilos["Cuerpo"], borderColor=cfg.COLOR_SECUNDARIO, backColor=cfg.COLOR_FONDO_PAGINA)
    for tipo_elemento, elemento in iterar_elementos_documento(documento_word):
        if tipo_elemento == "tabla":
            item_tabla = item_caja_tabla(elemento)
            if estado.agregar_contenido_a_bloque_activo(item_tabla, []):
                continue
            estado.cerrar_bloques_para_contenido_general()
            tabla_fuera_de_caja = tabla_docx_a_flujo(elemento, ancho_util, estilo_tabla_general)
            if tabla_fuera_de_caja is not None:
                estado.historia.extend([Spacer(1, 3), tabla_fuera_de_caja, Spacer(1, 3)])
            continue
        parrafo = elemento
        imagenes = extraer_imagenes_de_parrafo(parrafo, documento_word)
        if _parrafo_tiene_salto_de_pagina_previo(parrafo):
            estado.insertar_salto_de_pagina()
        if _parrafo_es_solo_salto_de_pagina(parrafo) and not imagenes:
            estado.insertar_salto_de_pagina()
            continue
        texto_html = parrafo_a_html(parrafo)
        texto_plano = parrafo.text or ""
        if estado.procesar_apertura_bloque_manual(texto_plano):
            continue
        if estado.procesar_fin_bloque_manual(texto_plano):
            continue
        if not texto_html.strip() and not imagenes:
            if estado.agregar_salto_a_bloque_activo():
                continue
            if estado.citas_pendientes:
                estado.vacios_desde_ultima_cita += 1
                if estado.vacios_desde_ultima_cita > estado.max_vacios_fusion_cita:
                    estado.vaciar_citas()
                    estado.vacios_desde_ultima_cita = 0
            if estado.consejos_pendientes:
                estado.vacios_desde_ultimo_consejo += 1
                if estado.vacios_desde_ultimo_consejo > estado.max_vacios_fusion_consejo:
                    estado.vaciar_consejos()
            if estado.infos_pendientes:
                estado.vacios_desde_ultima_info += 1
                if estado.vacios_desde_ultima_info > estado.max_vacios_fusion_info:
                    estado.vaciar_infos()
            estado.vaciar_lista()
            if not any([estado.dentro_info, estado.dentro_consejo_manual, estado.dentro_cita_manual, estado.dentro_npc_manual, estado.dentro_enemigo_manual, estado.dentro_aliado_manual, estado.dentro_tesoro_manual, estado.dentro_objeto_manual]):
                estado.historia.append(Spacer(1, 0))
            continue
        clave = estilo_para_parrafo(parrafo)
        consejo_por_estilo = clave == "CajaConsejoDm"
        consejo_por_prefijo = es_consejo_dm(texto_plano)
        info_por_estilo = clave == "CajaInfoAdicional"
        info_por_prefijo = es_info_adicional(texto_plano)
        es_lista = clave == "Lista" or (parrafo.style.name and "List Bullet" in parrafo.style.name) or es_lista_parrafo(parrafo)
        item_caja = item_caja_desde_parrafo(parrafo, texto_html)
        if estado.agregar_contenido_a_bloque_activo(item_caja if texto_html.strip() else None, imagenes):
            continue
        if imagenes and not texto_html.strip():
            if estado.consejos_pendientes:
                agregar_imagenes_a_bloque(estado.consejos_pendientes, imagenes)
                estado.vacios_desde_ultimo_consejo = 0
                continue
            if estado.infos_pendientes:
                agregar_imagenes_a_bloque(estado.infos_pendientes, imagenes)
                estado.vacios_desde_ultima_info = 0
                continue
            if estado.citas_pendientes:
                agregar_imagenes_a_bloque(estado.citas_pendientes, imagenes)
                estado.vacios_desde_ultima_cita = 0
                continue
        if clave not in ("CajaConsejoDm", "CajaInfoAdicional", "CajaCita"):
            bloque_embebido = _extraer_bloque_prefijo_embebido(texto_html, texto_plano, es_consejo_dm)
            tipo_bloque_embebido = "consejo"
            if bloque_embebido is None:
                bloque_embebido = _extraer_bloque_prefijo_embebido(texto_html, texto_plano, es_info_adicional)
                tipo_bloque_embebido = "info"
            if bloque_embebido is not None:
                antes_html, _, bloque_html, _ = bloque_embebido
                estado.vaciar_consejos(); estado.vaciar_infos(); estado.emitir_info_adicional(); estado.emitir_consejo_manual(); estado.emitir_cita_manual(); estado.vaciar_citas()
                if es_lista:
                    estado.items_lista_actual.append(antes_html)
                else:
                    estado.vaciar_lista(); estado.historia.append(Paragraph(antes_html, estilos[clave]))
                estado.vaciar_lista()
                item_bloque = item_caja_plano(bloque_html)
                if tipo_bloque_embebido == "consejo":
                    estado.consejos_pendientes.append(item_bloque)
                    if imagenes:
                        agregar_imagenes_a_bloque(estado.consejos_pendientes, imagenes)
                    estado.modo_consejos = "prefijo"
                    estado.vacios_desde_ultimo_consejo = 0
                else:
                    estado.infos_pendientes.append(item_bloque)
                    if imagenes:
                        agregar_imagenes_a_bloque(estado.infos_pendientes, imagenes)
                    estado.modo_infos = "prefijo"
                    estado.vacios_desde_ultima_info = 0
                continue
        if consejo_por_estilo or consejo_por_prefijo:
            nuevo_modo = "estilo" if consejo_por_estilo else "prefijo"
            if estado.consejos_pendientes and estado.modo_consejos != nuevo_modo:
                estado.vaciar_consejos()
            estado.vaciar_infos(); estado.emitir_consejo_manual(); estado.emitir_cita_manual(); estado.vaciar_citas(); estado.vaciar_lista()
            item_consejo = item_caja if consejo_por_estilo else item_caja_plano(texto_html)
            if texto_html.strip():
                estado.consejos_pendientes.append(item_consejo)
            if imagenes:
                agregar_imagenes_a_bloque(estado.consejos_pendientes, imagenes)
            estado.modo_consejos = nuevo_modo
            estado.vacios_desde_ultimo_consejo = 0
            continue
        if info_por_estilo or info_por_prefijo:
            nuevo_modo = "estilo" if info_por_estilo else "prefijo"
            if estado.infos_pendientes and estado.modo_infos != nuevo_modo:
                estado.vaciar_infos()
            estado.vaciar_consejos(); estado.vaciar_citas(); estado.vaciar_lista(); estado.emitir_cita_manual()
            if texto_html.strip():
                estado.infos_pendientes.append(item_caja)
            if imagenes:
                agregar_imagenes_a_bloque(estado.infos_pendientes, imagenes)
            estado.modo_infos = nuevo_modo
            estado.vacios_desde_ultima_info = 0
            continue
        if es_lista:
            if estado.consejos_pendientes and estado.modo_consejos != "prefijo":
                if texto_html.strip():
                    estado.consejos_pendientes.append(item_caja)
                if imagenes:
                    agregar_imagenes_a_bloque(estado.consejos_pendientes, imagenes)
                estado.vacios_desde_ultimo_consejo = 0
                continue
            if estado.infos_pendientes and estado.modo_infos != "prefijo":
                if texto_html.strip():
                    estado.infos_pendientes.append(item_caja)
                if imagenes:
                    agregar_imagenes_a_bloque(estado.infos_pendientes, imagenes)
                estado.vacios_desde_ultima_info = 0
                continue
            if estado.citas_pendientes:
                if texto_html.strip():
                    estado.citas_pendientes.append(item_caja)
                if imagenes:
                    agregar_imagenes_a_bloque(estado.citas_pendientes, imagenes)
                estado.vacios_desde_ultima_cita = 0
                continue
            estado.vaciar_consejos(); estado.vaciar_infos(); estado.emitir_consejo_manual(); estado.emitir_cita_manual(); estado.vaciar_citas()
            estado.items_lista_actual.append(texto_html)
            continue
        estado.vaciar_lista()
        if clave == "CajaCita":
            estado.vaciar_consejos(); estado.vaciar_infos(); estado.emitir_consejo_manual(); estado.emitir_cita_manual()
            if texto_html.strip():
                estado.citas_pendientes.append(item_caja)
            if imagenes:
                agregar_imagenes_a_bloque(estado.citas_pendientes, imagenes)
            estado.vacios_desde_ultima_cita = 0
            continue
        estado.vaciar_consejos(); estado.vaciar_infos(); estado.emitir_consejo_manual(); estado.emitir_cita_manual(); estado.vaciar_citas(); estado.vacios_desde_ultima_cita = 0
        if texto_html.strip():
            estado.historia.append(Paragraph(texto_html, estilos[clave]))
        if imagenes:
            for flujo_imagen in crear_flujos_imagenes(imagenes, ancho_util):
                if flujo_imagen is not None:
                    estado.historia.extend([Spacer(1, 3), flujo_imagen, Spacer(1, 3)])
    estado.cerrar_todo_al_final()
    documento_pdf.multiBuild(historia)
    print(f"✅ PDF generado: {ruta_pdf}")
